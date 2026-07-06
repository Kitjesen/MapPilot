from __future__ import annotations

import re
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "07-testing" / "field-runs"
MD_PATH = OUT_DIR / "LingTu_Navigation_Progress_Report_2026-07-06.md"
DOCX_PATH = OUT_DIR / "LingTu_Navigation_Progress_Report_2026-07-06.docx"
PDF_PATH = OUT_DIR / "LingTu_Navigation_Progress_Report_2026-07-06.pdf"
PREVIEW_DIR = OUT_DIR / "LingTu_Navigation_Progress_Report_2026-07-06_preview"

DOC_FONT = "Microsoft YaHei"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(30, 30, 30)
MUTED = RGBColor(90, 90, 90)
FILL = "F2F4F7"
BORDER = "C9D2DD"


def clean_inline(text: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", text)


def parse_table(lines: list[str], index: int) -> tuple[list[list[str]], int]:
    rows = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
        index += 1
    return rows, index


def table_has_separator(rows: list[list[str]]) -> bool:
    return (
        len(rows) > 1
        and all(set(cell.replace(":", "").replace(" ", "")) <= {"-"} for cell in rows[1])
    )


def docx_widths(headers: list[str]) -> list[float]:
    key = "|".join(headers)
    if headers == ["阶段", "目标", "要做的事", "验收标准"]:
        return [0.85, 1.45, 2.65, 1.55]
    if headers == ["序号", "已完成事项", "实际做了什么", "证据或结果"]:
        return [0.45, 1.25, 2.45, 2.35]
    if headers == ["序号", "未完成事项", "当前真实状态", "影响"]:
        return [0.45, 1.55, 2.15, 2.35]
    if headers == ["链路段", "输入", "输出", "当前状态"]:
        return [1.15, 1.75, 1.95, 1.65]
    if headers == ["证据", "结果", "解释"]:
        return [2.65, 0.85, 3.0]
    if headers == ["问题", "当前观察", "可能原因", "下一步验证"]:
        return [1.5, 1.55, 2.3, 1.15]
    if headers == ["优先级", "任务", "产出物", "是否阻塞后续"]:
        return [0.7, 2.25, 2.4, 1.15]
    if headers == ["风险", "为什么重要", "控制方式"]:
        return [1.7, 2.4, 2.4]
    if headers == ["文件", "作用"]:
        return [3.15, 3.35]
    if headers == ["判断", "结论"]:
        return [1.45, 5.05]
    col_count = len(headers)
    if col_count == 4:
        return [1.15, 2.1, 1.0, 2.15]
    if col_count == 3:
        return [2.1, 1.4, 3.0]
    if col_count == 2:
        return [1.6, 4.9]
    return [6.5 / col_count] * col_count


def pdf_widths(headers: list[str]) -> list[float]:
    if headers == ["阶段", "目标", "要做的事", "验收标准"]:
        return [0.75, 1.3, 2.35, 1.6]
    if headers == ["序号", "已完成事项", "实际做了什么", "证据或结果"]:
        return [0.45, 1.2, 2.25, 2.1]
    if headers == ["序号", "未完成事项", "当前真实状态", "影响"]:
        return [0.45, 1.4, 1.95, 2.2]
    if headers == ["链路段", "输入", "输出", "当前状态"]:
        return [1.05, 1.65, 1.75, 1.55]
    if headers == ["证据", "结果", "解释"]:
        return [2.5, 0.8, 2.7]
    if headers == ["问题", "当前观察", "可能原因", "下一步验证"]:
        return [1.25, 1.45, 2.2, 1.1]
    if headers == ["优先级", "任务", "产出物", "是否阻塞后续"]:
        return [0.65, 2.05, 2.2, 1.1]
    if headers == ["风险", "为什么重要", "控制方式"]:
        return [1.55, 2.2, 2.25]
    if headers == ["文件", "作用"]:
        return [3.0, 3.0]
    if headers == ["判断", "结论"]:
        return [1.3, 4.7]
    col_count = len(headers)
    if col_count == 4:
        return [1.15, 2.0, 0.9, 1.95]
    if col_count == 3:
        return [2.0, 1.3, 2.7]
    if col_count == 2:
        return [1.6, 4.4]
    return [6.0 / col_count] * col_count


def set_run_font(run, size=None, bold=None, color=None, font=DOC_FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_spacing(paragraph, before=0, after=6, line=1.10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade_cell(cell, fill: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches: float):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def add_docx_table(doc: Document, rows: list[list[str]]):
    headers = rows[0]
    body = rows[2:] if table_has_separator(rows) else rows[1:]
    col_count = len(headers)
    widths = docx_widths(headers)
    table = doc.add_table(rows=1, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell)
        shade_cell(cell, FILL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        set_spacing(p, after=0, line=1.05)
        r = p.add_run(clean_inline(header))
        set_run_font(r, 9.5, bold=True, color=INK)
    for row in body:
        cells = table.add_row().cells
        row = row + [""] * (col_count - len(row))
        for idx, value in enumerate(row[:col_count]):
            cell = cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            set_spacing(p, after=0, line=1.05)
            r = p.add_run(clean_inline(value))
            set_run_font(r, 9, color=INK)
    doc.add_paragraph()


def build_docx(lines: list[str]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = DOC_FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = DOC_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    set_spacing(header, after=0)
    run = header.add_run("LingTu Navigation Validation Report")
    set_run_font(run, 9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(footer, after=0)
    run = footer.add_run("2026-07-06 | sunrise native DDS")
    set_run_font(run, 9, color=MUTED)

    index = 0
    while index < len(lines):
        text = lines[index].strip()
        if not text:
            index += 1
            continue
        if text.startswith("```"):
            code = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index])
                index += 1
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            set_table_borders(table, color="D7DCE3")
            cell = table.cell(0, 0)
            shade_cell(cell, "F7F8FA")
            set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
            p = cell.paragraphs[0]
            set_spacing(p, after=0, line=1.05)
            run = p.add_run("\n".join(code))
            set_run_font(run, 9, color=RGBColor(40, 40, 40), font="Consolas")
            doc.add_paragraph()
            index += 1
            continue
        if text.startswith("|"):
            rows, index = parse_table(lines, index)
            add_docx_table(doc, rows)
            continue
        if text.startswith("# "):
            p = doc.add_paragraph()
            set_spacing(p, after=8)
            run = p.add_run(text[2:])
            set_run_font(run, 22, bold=True, color=RGBColor(0, 0, 0))
        elif text.startswith("## "):
            doc.add_paragraph(text[3:], style="Heading 1")
        elif text.startswith("### "):
            doc.add_paragraph(text[4:], style="Heading 2")
        elif text.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_spacing(p, after=4, line=1.167)
            run = p.add_run(clean_inline(text[2:]))
            set_run_font(run, 10.5, color=INK)
        elif re.match(r"^\d+\. ", text):
            p = doc.add_paragraph(style="List Number")
            set_spacing(p, after=4, line=1.167)
            run = p.add_run(clean_inline(re.sub(r"^\d+\. ", "", text)))
            set_run_font(run, 10.5, color=INK)
        else:
            p = doc.add_paragraph()
            set_spacing(p, after=6)
            run = p.add_run(clean_inline(text))
            set_run_font(run, 11, color=INK)
        index += 1
    doc.save(DOCX_PATH)


def para_text(text: str) -> str:
    text = escape(text)
    return re.sub(r"`([^`]+)`", r"<font name='Courier'>\1</font>", text)


def add_pdf_table(story, styles, rows: list[list[str]]):
    headers = rows[0]
    body = rows[2:] if table_has_separator(rows) else rows[1:]
    col_count = len(headers)
    widths = pdf_widths(headers)
    data = [[Paragraph(para_text(h), styles["TableHeadCN"]) for h in headers]]
    for row in body:
        row = row + [""] * (col_count - len(row))
        data.append([Paragraph(para_text(c), styles["TableCN"]) for c in row[:col_count]])
    table = Table(data, colWidths=[width * inch for width in widths], repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C9D2DD")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 8))


def build_pdf(lines: list[str]) -> int:
    pdfmetrics.registerFont(TTFont("NotoSC", r"C:\Windows\Fonts\NotoSansSC-VF.ttf"))
    pdfmetrics.registerFont(TTFont("NotoSC-Bold", r"C:\Windows\Fonts\simhei.ttf"))
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("TitleCN", fontName="NotoSC-Bold", fontSize=20, leading=25, textColor=colors.black, spaceAfter=10))
    styles.add(ParagraphStyle("H1CN", fontName="NotoSC-Bold", fontSize=15, leading=20, textColor=colors.HexColor("#2E74B5"), spaceBefore=14, spaceAfter=7, keepWithNext=True))
    styles.add(ParagraphStyle("H2CN", fontName="NotoSC-Bold", fontSize=12, leading=16, textColor=colors.HexColor("#2E74B5"), spaceBefore=9, spaceAfter=5, keepWithNext=True))
    styles.add(ParagraphStyle("BodyCN", fontName="NotoSC", fontSize=10.4, leading=15.5, alignment=TA_LEFT, spaceAfter=6))
    styles.add(ParagraphStyle("BulletCN", fontName="NotoSC", fontSize=10.2, leading=15, leftIndent=16, firstLineIndent=-8, spaceAfter=4))
    styles.add(ParagraphStyle("TableCN", fontName="NotoSC", fontSize=8.2, leading=11.2))
    styles.add(ParagraphStyle("TableHeadCN", fontName="NotoSC-Bold", fontSize=8.2, leading=11.2))
    styles.add(ParagraphStyle("CodeCN", fontName="Courier", fontSize=8, leading=10.2, textColor=colors.HexColor("#333333"), backColor=colors.HexColor("#F7F8FA"), leftIndent=4, rightIndent=4, spaceAfter=8))
    story = []

    index = 0
    while index < len(lines):
        text = lines[index].strip()
        if not text:
            index += 1
            continue
        if text.startswith("```"):
            code = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index])
                index += 1
            story.append(Preformatted("\n".join(code), styles["CodeCN"]))
            story.append(Spacer(1, 6))
            index += 1
            continue
        if text.startswith("|"):
            rows, index = parse_table(lines, index)
            add_pdf_table(story, styles, rows)
            continue
        if text.startswith("# "):
            story.append(Paragraph(para_text(text[2:]), styles["TitleCN"]))
        elif text.startswith("## "):
            story.append(Paragraph(para_text(text[3:]), styles["H1CN"]))
        elif text.startswith("### "):
            story.append(Paragraph(para_text(text[4:]), styles["H2CN"]))
        elif text.startswith("- "):
            story.append(Paragraph("• " + para_text(text[2:]), styles["BulletCN"]))
        elif re.match(r"^\d+\. ", text):
            story.append(Paragraph("• " + para_text(re.sub(r"^\d+\. ", "", text)), styles["BulletCN"]))
        else:
            story.append(Paragraph(para_text(text), styles["BodyCN"]))
        index += 1

    def footer_canvas(canvas, doc):
        canvas.saveState()
        canvas.setFont("NotoSC", 8)
        canvas.setFillColor(colors.HexColor("#666666"))
        canvas.drawString(doc.leftMargin, 0.45 * inch, "LingTu Navigation Validation Report")
        canvas.drawRightString(letter[0] - doc.rightMargin, 0.45 * inch, f"Page {doc.page}")
        canvas.restoreState()

    pdf_doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
        topMargin=0.85 * inch,
        bottomMargin=0.75 * inch,
    )
    pdf_doc.build(story, onFirstPage=footer_canvas, onLaterPages=footer_canvas)
    return len(story)


def render_pdf_preview() -> int:
    import pypdfium2 as pdfium
    from PIL import Image, ImageDraw

    PREVIEW_DIR.mkdir(exist_ok=True)
    for path in PREVIEW_DIR.glob("page-*.png"):
        path.unlink()
    for path in PREVIEW_DIR.glob("contact-sheet.png"):
        path.unlink()

    pdf = pdfium.PdfDocument(str(PDF_PATH))
    image_paths = []
    for index in range(len(pdf)):
        image = pdf[index].render(scale=1.6).to_pil().convert("RGB")
        out = PREVIEW_DIR / f"page-{index + 1:02d}.png"
        image.save(out)
        image_paths.append(out)

    thumbs = []
    for path in image_paths:
        image = Image.open(path).convert("RGB")
        image.thumbnail((360, 480))
        canvas = Image.new("RGB", (390, 530), "white")
        canvas.paste(image, ((390 - image.width) // 2, 20))
        draw = ImageDraw.Draw(canvas)
        draw.text((15, 500), path.stem, fill=(0, 0, 0))
        thumbs.append(canvas)
    cols = 2
    rows = (len(thumbs) + cols - 1) // cols
    sheet = Image.new("RGB", (cols * 390, rows * 530), (245, 245, 245))
    for index, thumb in enumerate(thumbs):
        sheet.paste(thumb, ((index % cols) * 390, (index // cols) * 530))
    sheet.save(PREVIEW_DIR / "contact-sheet.png")
    return len(pdf)


def main() -> None:
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    build_docx(lines)
    build_pdf(lines)
    page_count = render_pdf_preview()
    print(DOCX_PATH)
    print(PDF_PATH)
    print(MD_PATH)
    print(PREVIEW_DIR / "contact-sheet.png")
    print(f"pages={page_count}")


if __name__ == "__main__":
    main()
